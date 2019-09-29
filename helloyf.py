import hashlib
import json
import random
import time
from urllib import request

import requests
from bs4 import BeautifulSoup

import docx
from docx.shared import RGBColor
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_COLOR

# 百度api调用
appid = 'yf作业还能再多点吗'
appkey = '我好难啊'
api_url = 'http://api.fanyi.baidu.com/api/trans/vip/translate'

article_wordlist = []
word_trans_dict = dict() # 为生词创建英文 -> 中文的映射字典

# 提取介词短语
prep = set(('on','in','to','of','up','down','by','with','about','without','into','from','away','off'))
lc = set(('and', 'or'))
bd = set((',','.',';','\''))
dc = set(('it','I','you','they','he','she','it'))
fin = open('article.txt', 'r')
fout = open('terms.txt', 'w')

def spider():
    wordlist = set()
    fin = open('article.txt', 'r')
    article_text = fin.read()
    url = 'https://www.lextutor.ca/cgi-bin/vp/comp/output.pl'

    headers = {
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'
        'AppleWebKit/537.36 (KHTML, like Gecko)'
        'Chrome/77.0.3865.90 Safari/537.36',
        'origin': 'https://www.lextutor.ca',
        'referer': 'https://www.lextutor.ca/vp/comp/'
    }

    data = {
        'vintage': 'bnc_coca',
        'via': 'formulaire',
        'text_name': 'Untitled',
        'wants_edit': 'on',
        'text_input': article_text,
        'prop_handle': 'ignore',
        'compounds': 'false'
    }

    reponse = requests.post(url, data=data, headers=headers)
    print('Response Status Code', reponse.status_code)

    if reponse.status_code == 200:
        html = reponse.text
        bsobj = BeautifulSoup(html, 'lxml')
        for idx in range(5, 26):
            liststr = 'list_' + str(idx)
            district_tags = bsobj.find_all('div',attrs={'id': liststr})
            # print(liststr+'\n')
            # print(district_tags[0].get_text())
            tem_wordlist = district_tags[0].get_text().split(' ')
            for tem_word in tem_wordlist:
                if tem_word == '':
                    pass
                else:
                    left = tem_word.find('[')-1
                    right = tem_word.find(']')
                    tem_word = tem_word[:left]
                    wordlist.add(tem_word)

    return wordlist
    fin.close()

def random_number():
    return str(random.randint(0, 99999))

def generate_md5(init_str):
    m = hashlib .md5()
    m.update(init_str.encode('utf-8'))
    return m.hexdigest()

# 将生词上色
def drawcolor(_wordlist):
    list_strange_word = []
    list_prep_term = []
    document = Document() # 新建word
    p = document.add_paragraph()
    for idx in range(len(article_wordlist)):
        if (article_wordlist[idx] in _wordlist) == True:
            run = p.add_run(article_wordlist[idx] +'('+ word_trans_dict[article_wordlist[idx]]+')')
            list_strange_word.append(article_wordlist[idx])
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif (article_wordlist[idx] in prep) == True and \
            (article_wordlist[idx-1][-1:] in bd) == False and \
            (article_wordlist[idx-1] in prep) == False and \
            (article_wordlist[idx-1] in lc) == False:
            run = p.add_run(article_wordlist[idx-1]+' '+article_wordlist[idx])
            run.font.highlight_color = WD_COLOR_INDEX.PINK
            run.underline = True
            list_prep_term.append(article_wordlist[idx-1]+' '+article_wordlist[idx])
        elif (article_wordlist[idx] in prep) == True and \
            (article_wordlist[idx-1][-2:] == 'ly'):
            run = p.add_run(article_wordlist[idx-2]+' '+article_wordlist[idx-1]+' '+article_wordlist[idx])
            run.font.highlight_color = WD_COLOR_INDEX.PINK
            run.underline = True
            list_prep_term.append(article_wordlist[idx-2]+' '+article_wordlist[idx-1]+' '+article_wordlist[idx])
        else:
            run = p.add_run(article_wordlist[idx])
            run.font.highlight_color = None
        run = p.add_run(' ')
    p1 = document.add_paragraph('障碍词块：')
    p1.bold = True
    for _item in list_strange_word:
        p1 = document.add_paragraph(_item+' '+word_trans_dict[_item])
    p1 = document.add_paragraph('写作词块：')
    p1.bold = True
    for _item in list_prep_term:
        p1 = document.add_paragraph(_item)
    document.save('test_out.docx')


# 寻找写作短语，将介词短语上色
def search_terms():
    pass

# 将生词和短语追加到阅读文件末尾
def file_cat():
    pass

def get_article_wordlist():
    f = open('article.txt', 'r')
    lines = f.readlines()

    for line in lines:
        print(line)
        splits = line.split(' ')
        for _word in splits:
            article_wordlist.append(_word)

if __name__ == '__main__':
    wl = spider() # 爬取单词列表
    salt = random_number() # 生成随机数
    for word in wl:
        params_str = appid+word+str(salt)+appkey
        jmstr = generate_md5(params_str)
        newurl = api_url + '?' + \
                'q=' + word + '&' + \
                'from=en&to=zh&' + \
                'appid=' + appid + '&' + \
                'salt=' + salt + '&' + \
                'sign=' + jmstr
        
        res = requests.get(newurl)
        
        time.sleep(1)
        # print(len(res.text))
        # print(json.loads(res.text))
        zhword = json.loads(res.text)['trans_result'][0]['dst']
        word_trans_dict.update({word: zhword})
        print(word, zhword)
        # print(res.content['trans_result'][0]['dst'].decode())

    print(word_trans_dict)
    get_article_wordlist()
    # print(article_wordlist)
    drawcolor(wl)